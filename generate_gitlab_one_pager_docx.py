from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _set_compact_paragraph(p) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _set_compact_paragraph(p)
    run = p.add_run(text)
    run.font.size = Pt(10)


def _add_role(doc: Document, header: str, bullets: list[str]) -> None:
    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    run = p.add_run(header)
    run.bold = True
    run.font.size = Pt(10)
    for b in bullets:
        _add_bullet(doc, b)


def build_doc(out_path: Path) -> None:
    doc = Document()

    # Page setup (aim for true one-pager).
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # Header
    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UMAPRASATH UDAIYAR")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VP Engineering • Architecture • Transformation • AI-Driven Productivity")
    r.bold = True
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("+1 980-401-2087 • umaprasath.lru@gmail.com • linkedin.com/in/umaprasathudaiyar • Remote (US)")
    r.font.size = Pt(10)

    # Summary
    _add_heading(doc, "Executive Summary")
    for b in [
        "Senior technology executive with 23+ years leading enterprise-scale engineering across financial services, wealth management, consumer lending, e-commerce, and digital health.",
        "Built and scaled distributed orgs of 45–110+ engineers with full budget ownership ($3M–$5M), aligning architecture and execution to measurable business outcomes.",
        "Modernized architectures (monolith → service-oriented), raised quality/reliability, and improved delivery throughput via DevSecOps automation and AI-enabled workflows (including agentic automation and Claude Code/GitLab Duo-class assistants).",
    ]:
        _add_bullet(doc, b)

    # Strengths
    _add_heading(doc, "Signature strengths (fit for GitLab ELT pillar)")
    for b in [
        "Horizontal execution & engineering excellence: standards, governance, reusable primitives, and embedded teams that unblock verticals.",
        "Architecture modernization: domain decomposition, API/platform primitives, event-driven patterns, and cloud-native scalability for enterprise customers.",
        "World-scale operations: SLO-driven reliability, observability, incident learning loops, and compliance-first delivery in regulated environments.",
        "AI-assisted productivity: applied AI/automation to reduce non-creative work and accelerate teams; operationalized LLM-assisted development workflows to speed refactoring and review cycles.",
        "Cost & COGS discipline: cloud and vendor optimization without compromising availability or security.",
    ]:
        _add_bullet(doc, b)

    # Experience
    _add_heading(doc, "Professional Experience")
    _add_role(
        doc,
        "LendingPoint — Senior Vice President, Architecture & Applications | Irving, TX (Remote) | Mar 2025 – Dec 2025",
        [
            "Reported to CTO; owned $5M budget and led 45+ technologists across enterprise architecture and application development for a $500M+ lending platform.",
            "Set target architecture and execution strategy for core financial systems (payments, ledger, reconciliation), driving cloud modernization to improve reliability and reduce operational cost.",
            "Established architecture governance and engineering standards; evaluated AI/ML opportunities and developer productivity tooling to sustain competitive advantage.",
        ],
    )
    _add_role(
        doc,
        "Fidelity Investments — Director of Software Engineering | Salt Lake City, UT | Aug 2023 – Mar 2025",
        [
            "Led 60+ engineering leads, software engineers, and QA across Money Movement, Customer Management, and Account Management; owned $5M budget and AWS roadmap.",
            "Delivered secure, fault-tolerant payment and transfer capabilities integrating banking/payment APIs (including Plaid) and ledger systems for ACH, internal transfers, and wire settlements.",
            "Influenced cross-product architecture and compliance via Architecture Review Forum governance, improving consistency of design decisions across a matrixed organization.",
        ],
    )
    _add_role(
        doc,
        "My Muscle Chef — Head of Engineering | Sydney, Australia | Sep 2020 – Nov 2022",
        [
            "Executive leadership for a $350M B2C/B2B e-commerce business; managed $3M technology budget and a 45+ engineering organization.",
            "Led platform modernization to a headless architecture; improved conversion by 4% and increased incremental annual revenue through faster experimentation and reliability.",
            "Negotiated enterprise partnerships and rationalized vendors, reducing costs by 15% while accelerating roadmap delivery.",
        ],
    )
    _add_role(
        doc,
        "Healthdirect Australia — Engineering Manager | Sydney, Australia | Dec 2016 – Sep 2020",
        [
            "Led 3 engineering teams delivering API platforms serving 5M+ monthly transactions for government healthcare programs.",
            "Delivered compliant, mission-critical API infrastructure achieving 99.95% uptime; built talent programs resulting in 8 senior-level promotions and improved retention.",
        ],
    )

    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    r = p.add_run(
        "Earlier: Tata Consultancy Services (Delivery Manager / Architect / Centre of Excellence, APAC) • Hewlett Packard Singapore (Project Lead) • emeDigital Australia (Head of Digital)"
    )
    r.font.size = Pt(10)

    # Innovation
    _add_heading(doc, "Innovation (AI + Automation)")
    _add_bullet(
        doc,
        "AI in Lending: built an AI-powered credit decisioning and loan origination platform integrating predictive analytics, risk scoring, and human-in-the-loop automation (lending-ai.trotlabs.com).",
    )
    _add_bullet(
        doc,
        "Agentic AI in Support Ops: implemented AI-driven workflow automation (Vonage AI) to optimize customer support operations and reduce cost-to-serve.",
    )

    # Education
    _add_heading(doc, "Education")
    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    r = p.add_run("Stanford University Graduate School of Business — LEAD Certificate (Business Administration & Management)")
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    r = p.add_run("University of Madras — Bachelor of Engineering, Computer Science")
    r.font.size = Pt(10)

    # Selected tech (compact)
    _add_heading(doc, "Technology & Operating Model (select)")
    p = doc.add_paragraph()
    _set_compact_paragraph(p)
    r = p.add_run(
        "Cloud: AWS, Azure, Oracle Cloud Infrastructure • Platforms: Kubernetes, Kafka • Data: Postgres, MongoDB, DynamoDB • Practices: DevSecOps, CI/CD, Observability, IAM/CIAM, Architecture Governance"
    )
    r.font.size = Pt(10)

    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate GitLab-targeted one-page resume (DOCX).")
    parser.add_argument(
        "--out",
        default="Umaprasath_Udaiyar_GitLab_VP_Engineering_One_Pager.docx",
        help="Output DOCX path",
    )
    args = parser.parse_args()

    out_path = Path(args.out)
    build_doc(out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

